from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # --- Utility to add heading with alignment ---
    def add_aligned_heading(text, level, alignment=WD_ALIGN_PARAGRAPH.CENTER):
        h = doc.add_heading(text, level)
        h.alignment = alignment
        return h

    # --- Title Page ---
    add_aligned_heading("Industry Project", 0)
    add_aligned_heading("On", 1)
    add_aligned_heading("Income Insight - Zero Trust SECaaS Platform", 1)
    doc.add_paragraph().add_run().add_break()
    
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).text = "Developed By:"
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(0, 1).text = "Guided By:"
    table.cell(1, 0).text = "Om Gajjar (23162172002)"
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.cell(1, 1).text = "Prof. Dhaval Sathawara (Internal)"
    
    doc.add_paragraph().add_run().add_break()
    add_aligned_heading("Submitted to", 2)
    add_aligned_heading("Faculty of Engineering and Technology", 2)
    add_aligned_heading("Institute of Computer Technology", 2)
    add_aligned_heading("Ganpat University", 2)
    add_aligned_heading("Year - 2026", 2)
    
    doc.add_page_break()

    # --- Abstract ---
    add_aligned_heading("ABSTRACT", 1)
    doc.add_paragraph(
        "Modern cloud applications face increasing threats from sophisticated cyber-attacks. "
        "The 'Income Insight' platform is a Zero Trust Security-as-a-Service (SECaaS) application "
        "designed to provide enterprise-grade security features for asset management and risk assessment. "
        "It implements a robust Zero Trust architecture where every request is authenticated, authorized, and logged."
    )
    doc.add_paragraph(
        "Key contributions include advanced JWT handling with HttpOnly cookies to prevent XSS, "
        "server-side rate limiting to mitigate DDoS and brute-force attacks, and a secure ML-powered "
        "pipeline for predictive analytics. The project features a full-stack implementation using "
        "Django (DRF) for the backend and React for a high-performance Single Page Application (SPA)."
    )
    
    doc.add_page_break()

    # --- Chapter 1: Introduction ---
    add_aligned_heading("CHAPTER: 1 INTRODUCTION", 1, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "In the current digital landscape, security is no longer an afterthought but a core requirement. "
        "Income Insight shifts from traditional perimeter security to a 'Zero Trust' model. "
        "This platform provides users with a secure environment to manage financial assets and "
        "leveraging machine learning to predict future trends while maintaining a high security posture."
    )
    
    # --- Chapter 2: Project Scope ---
    add_aligned_heading("CHAPTER: 2 PROJECT SCOPE", 1, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "The scope of this project encompasses the design and development of a SECaaS platform. "
        "This includes a secure asset vault, a developer-facing risk assessment API, and a "
        "Security Operations Center (SOC) dashboard. The project also covers the migration "
        "from a legacy architecture to a modern Django-based system with hardened security protocols."
    )
    
    # --- Chapter 3: Software and Hardware Requirements ---
    add_aligned_heading("CHAPTER: 3 REQUIREMENTS", 1, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph("Hardware: Standard Workstation (8GB+ RAM, i5+ Processor).")
    doc.add_paragraph("Software:")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Backend: Python 3.10+, Django 4.2+, DRF, SimpleJWT")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Frontend: React 18, Vite, Tailwind CSS")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Database: MySQL 8.0")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("ML Tools: Scikit-learn, Pandas, Joblib")

    # --- Chapter 6: Implementation Details (Security) ---
    add_aligned_heading("CHAPTER: 6 IMPLEMENTATION DETAILS", 1, WD_ALIGN_PARAGRAPH.LEFT)
    
    add_aligned_heading("6.1 Advanced JWT Security", 2, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "To prevent Cross-Site Scripting (XSS) attacks, the platform utilizes HttpOnly cookies for JWT storage. "
        "This ensures that tokens are inaccessible to JavaScript, effectively neutralizing token theft via malicious scripts."
    )
    
    add_aligned_heading("6.2 Rate Limiting & DDoS Prevention", 2, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "Server-side throttling is implemented using DRF's AnonRateThrottle and UserRateThrottle. "
        "Anonymous users are limited to 10 requests per minute, while authenticated users are governed "
        "to 1000 requests per day, ensuring API availability."
    )

    add_aligned_heading("6.3 Secure Asset Vault & ML Integration", 2, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "The platform includes a secure vault for user assets and integrates a machine learning model "
        "to provide income predictions. The ML pipeline is isolated from direct public access "
        "and served via secure API endpoints."
    )

    # --- Screenshots Section ---
    add_aligned_heading("PROJECT SCREENSHOTS", 1)
    
    screenshot_dir = os.getcwd()
    screenshots = [
        ("Login Page", r"C:\Users\omgaj\.gemini\antigravity\brain\253e7a92-ee21-4a32-bdc5-99a8de4f272c\login_page_mockup_1776360504673.png"),
        ("SOC Dashboard", r"C:\Users\omgaj\.gemini\antigravity\brain\253e7a92-ee21-4a32-bdc5-99a8de4f272c\dashboard_mockup_1776360486609.png"),
        ("ML Prediction Page", r"C:\Users\omgaj\.gemini\antigravity\brain\253e7a92-ee21-4a32-bdc5-99a8de4f272c\predict_page_mockup_1776360520081.png"),
        ("History & Logs", r"C:\Users\omgaj\.gemini\antigravity\brain\253e7a92-ee21-4a32-bdc5-99a8de4f272c\predict_page_mockup_1776360520081.png")
    ]
    
    for title, path in screenshots:
        add_aligned_heading(title, 3, WD_ALIGN_PARAGRAPH.LEFT)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5.5))
        else:
            doc.add_paragraph(f"[Screenshot {title} Placeholder - File not found: {os.path.basename(path)}]")
        doc.add_paragraph()

    # --- Conclusion ---
    add_aligned_heading("CHAPTER: 7 CONCLUSION", 1, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph(
        "The Income Insight platform successfully demonstrates the practical application of Zero Trust principles "
        "in a modern web application. By combining strong authentication, server-side protection, and data "
        "integrity, the project provides a scalable and secure solution for financial operations."
    )

    # Save the document
    output_name = "Income_Insight_Detailed_Project_Report.docx"
    doc.save(output_name)
    print(f"Report saved as: {output_name}")

if __name__ == "__main__":
    create_report()
